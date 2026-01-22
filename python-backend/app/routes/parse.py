"""Document parsing endpoints with multi-format support"""

from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from typing import Optional, Tuple, Dict, Any
import tempfile
import os
import re
import zipfile
import io

from app.config import settings

router = APIRouter()


@router.post("/parse-file")
async def parse_file(
    file: UploadFile = File(...),
    mimeType: str = Form(None)
):
    """Parse various file formats and extract text"""
    
    content = await file.read()
    mime = mimeType or file.content_type or ""
    filename = file.filename or "uploaded"
    
    text, metadata = await parse_content(content, mime, filename)
    
    # Clean up text
    text = re.sub(r'\s+', ' ', text).strip()
    
    print(f"✅ Parsed {len(text)} characters from {filename}")
    return {"text": text, "length": len(text), "metadata": metadata}


async def parse_content(content: bytes, mime: str, filename: str = "") -> Tuple[str, Dict[str, Any]]:
    """
    Parse content based on MIME type.
    Returns (text, metadata) tuple.
    """
    mime_lower = mime.lower()
    metadata = {"filename": filename}
    
    print(f"📄 Parsing: {filename} ({mime})")
    
    try:
        # PDF
        if "pdf" in mime_lower:
            text = await parse_pdf_with_pages(content)
            return text, metadata
        
        # Excel
        elif "spreadsheet" in mime_lower or "xlsx" in mime_lower or "xls" in mime_lower:
            text = await parse_excel(content, filename)
            return text, metadata
        
        # CSV
        elif "csv" in mime_lower:
            text = await parse_csv(content)
            return text, metadata
        
        # PowerPoint
        elif "presentation" in mime_lower or "pptx" in mime_lower:
            text = await parse_pptx(content, filename)
            return text, metadata
        
        # Word
        elif "wordprocessing" in mime_lower or "docx" in mime_lower:
            text = await parse_docx(content)
            return text, metadata
        
        # Images
        elif mime_lower.startswith("image/"):
            text = await parse_image(content, mime)
            return text, {"type": "image", **metadata}
        
        # Plain text / Markdown
        elif "text" in mime_lower:
            text = content.decode("utf-8", errors="ignore")
            return text, metadata
        
        # Default: try as text
        else:
            try:
                text = content.decode("utf-8", errors="ignore")
                return text, metadata
            except:
                return "[Unable to parse file format]", metadata
                
    except Exception as e:
        print(f"❌ Parse error: {str(e)}")
        return f"[Parse error: {str(e)}]", metadata


async def parse_pdf_with_pages(content: bytes) -> str:
    """Parse PDF with page number markers"""
    from pypdf import PdfReader
    
    try:
        reader = PdfReader(io.BytesIO(content))
        text_parts = []
        total_text = ""
        
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(f"[PAGE {page_num}]\n{page_text}")
                total_text += page_text
        
        # Check for scanned PDF (minimal text)
        if len(total_text.strip()) < 100:
            print("📄 PDF appears scanned, attempting Gemini Vision...")
            return await parse_scanned_pdf_with_vision(content)
        
        return "\n\n".join(text_parts)
        
    except Exception as e:
        return f"[PDF parsing failed: {str(e)}]"


async def parse_scanned_pdf_with_vision(content: bytes) -> str:
    """Use Gemini Vision to extract text from scanned PDF"""
    try:
        import google.generativeai as genai
        from PIL import Image
        import pdf2image
        
        # This requires poppler - fallback to simple message if not available
        try:
            images = pdf2image.convert_from_bytes(content, first_page=1, last_page=5)
        except Exception:
            return "[Scanned PDF detected. Install poppler for OCR support.]"
        
        genai.configure(api_key=settings.GOOGLE_API_KEY)
        model = genai.GenerativeModel("gemini-2.5-flash-lite")
        
        all_text = []
        for i, img in enumerate(images[:5], 1):  # Limit to 5 pages
            # Convert PIL to bytes
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format='PNG')
            img_bytes = img_byte_arr.getvalue()
            
            response = model.generate_content([
                "Extract all text from this document page. Preserve the structure. If there are tables, format them clearly.",
                {"mime_type": "image/png", "data": img_bytes}
            ])
            
            if response.text:
                all_text.append(f"[PAGE {i}]\n{response.text}")
        
        return "\n\n".join(all_text) if all_text else "[Could not extract text from scanned PDF]"
        
    except Exception as e:
        return f"[Scanned PDF vision parsing failed: {str(e)}]"


async def parse_excel(content: bytes, filename: str = "") -> str:
    """Parse Excel using pandas - Markdown for small, Row-based for large"""
    import pandas as pd
    
    try:
        # Read Excel
        excel_file = io.BytesIO(content)
        
        # Check all sheets
        xl = pd.ExcelFile(excel_file)
        all_text = []
        
        for sheet_name in xl.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            
            if df.empty:
                continue
            
            all_text.append(f"## Sheet: {sheet_name}")
            
            # Small/Medium: Markdown table (< 50 rows)
            if len(df) <= 50:
                try:
                    from tabulate import tabulate
                    md_table = tabulate(df, headers='keys', tablefmt='pipe', showindex=False)
                    all_text.append(md_table)
                except:
                    all_text.append(df.to_string())
            
            # Large: Row-based extraction
            else:
                columns = df.columns.tolist()
                for idx, row in df.iterrows():
                    row_text = ", ".join([f"{col}: {row[col]}" for col in columns if pd.notna(row[col])])
                    if row_text:
                        all_text.append(f"Row {idx + 1}: {row_text}")
        
        result = "\n\n".join(all_text)
        return result if result else "[Excel file appears empty]"
        
    except Exception as e:
        return f"[Excel parsing failed: {str(e)}]"


async def parse_csv(content: bytes) -> str:
    """Parse CSV using pandas"""
    import pandas as pd
    
    try:
        df = pd.read_csv(io.BytesIO(content))
        
        if df.empty:
            return "[CSV file is empty]"
        
        # Small: Markdown
        if len(df) <= 50:
            try:
                from tabulate import tabulate
                return tabulate(df, headers='keys', tablefmt='pipe', showindex=False)
            except:
                return df.to_string()
        
        # Large: Row-based
        columns = df.columns.tolist()
        rows = []
        for idx, row in df.iterrows():
            row_text = ", ".join([f"{col}: {row[col]}" for col in columns if pd.notna(row[col])])
            if row_text:
                rows.append(f"Row {idx + 1}: {row_text}")
        
        return "\n".join(rows)
        
    except Exception as e:
        return f"[CSV parsing failed: {str(e)}]"


async def parse_image(content: bytes, mime_type: str) -> str:
    """Parse image using Gemini Vision"""
    try:
        import google.generativeai as genai
        
        genai.configure(api_key=settings.GOOGLE_API_KEY)
        model = genai.GenerativeModel("gemini-2.5-flash-lite")
        
        # Determine image type
        img_type = mime_type.split("/")[-1]
        if img_type == "jpg":
            img_type = "jpeg"
        
        response = model.generate_content([
            """Analyze this image and provide a detailed description:
            1. If it contains text, extract ALL the text.
            2. If it's a chart/diagram, describe what it shows with data points.
            3. If it's a photo, describe the contents in detail.
            Be thorough - this description will be used for search.""",
            {"mime_type": mime_type, "data": content}
        ])
        
        if response.text:
            return f"[IMAGE DESCRIPTION]\n{response.text}"
        
        return "[Could not analyze image]"
        
    except Exception as e:
        return f"[Image parsing failed: {str(e)}]"


async def parse_pptx(content: bytes, filename: str = "") -> str:
    """Parse PPTX using zipfile and XML parsing"""
    try:
        text_parts = [f"[File Name: {filename}]"] if filename else []
        
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            slide_files = sorted([
                f for f in zf.namelist() 
                if re.match(r'ppt/slides/slide\d+\.xml$', f)
            ], key=lambda x: int(re.search(r'slide(\d+)', x).group(1)))
            
            for slide_num, slide_file in enumerate(slide_files, 1):
                xml_content = zf.read(slide_file).decode('utf-8')
                matches = re.findall(r'<a:t>([^<]*)</a:t>', xml_content)
                slide_text = ' '.join([m.strip() for m in matches if m.strip()])
                if slide_text:
                    text_parts.append(f"[SLIDE {slide_num}]\n{slide_text}")
            
            # Notes
            notes_files = [f for f in zf.namelist() 
                         if re.match(r'ppt/notesSlides/notesSlide\d+\.xml$', f)]
            
            for notes_file in notes_files:
                xml_content = zf.read(notes_file).decode('utf-8')
                matches = re.findall(r'<a:t>([^<]*)</a:t>', xml_content)
                notes_text = ' '.join([m.strip() for m in matches if m.strip()])
                if notes_text:
                    text_parts.append(f"[NOTES]\n{notes_text}")
        
        return '\n\n'.join(text_parts) if text_parts else "[PPTX appears empty or image-only]"
        
    except Exception as e:
        return f"[PPTX parsing failed: {str(e)}]"


async def parse_docx(content: bytes) -> str:
    """Parse DOCX using python-docx"""
    from docx import Document
    
    try:
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                # Check if it's a heading
                if para.style and 'Heading' in para.style.name:
                    text_parts.append(f"\n## {para.text}\n")
                else:
                    text_parts.append(para.text)
        
        # Also extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                table_text.append(row_text)
            if table_text:
                text_parts.append("\n" + "\n".join(table_text) + "\n")
        
        return '\n'.join(text_parts)
        
    except Exception as e:
        return f"[DOCX parsing failed: {str(e)}]"

